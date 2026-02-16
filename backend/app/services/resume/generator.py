from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants
import os

class ResumeGenerator:
    def __init__(self):
        pass

    def add_hyperlink(self, paragraph, url, text, color="0000FF", underline=True, bold=False):
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id, )

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)

        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink

    def generate(self, data: dict, output_path: str):
        document = Document()
        
        # Margins
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)

        # Header
        self._add_header(document, data)
        self._add_summary(document, data.get('summary', ''))
        self._add_education(document, data.get('education', []))
        self._add_certifications(document, data.get('certifications', []))
        self._add_experience(document, data.get('experience', []))
        self._add_projects(document, data.get('projects', []))
        self._add_skills(document, data.get('skills', []))

        document.save(output_path)
        return output_path

    def _add_section_title(self, doc, title):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = False
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.text = title.upper()
        
        p_element = p._p
        pPr = p_element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4') 
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        p.paragraph_format.space_after = Pt(3)

    def _add_header(self, doc, data):
        name_paragraph = doc.add_paragraph()
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_paragraph.paragraph_format.space_after = Pt(2)
        
        name_run = name_paragraph.add_run(data.get('name', 'NAME'))
        name_run.bold = False 
        name_run.font.size = Pt(20) 
        name_run.font.name = 'Times New Roman'

        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.paragraph_format.space_after = Pt(6)
        
        run = contact_paragraph.add_run(data.get('phone', ''))
        run.font.size = Pt(9)
        
        contact_paragraph.add_run('  |  ')
        
        email = data.get('email', '')
        self.add_hyperlink(contact_paragraph, f'mailto:{email}', email, "000000", False)
        
        contact_paragraph.add_run('  |  ')
        
        linkedin = data.get('linkedin', '')
        self.add_hyperlink(contact_paragraph, linkedin, 'LinkedIn', "0000FF", True)

    def _add_summary(self, doc, summary):
        if not summary: return
        self._add_section_title(doc, 'Professional Summary')
        p = doc.add_paragraph(summary)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)

    def _add_education(self, doc, education_list):
        if not education_list: return
        self._add_section_title(doc, 'Education')
        for edu in education_list:
            self._add_row_two_cols(doc, edu['school'], True, False, edu['location'], False, False)
            self._add_row_two_cols(doc, edu['degree'], False, True, edu['date'], False, False, space_after=Pt(2))
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _add_certifications(self, doc, certs):
        if not certs: return
        self._add_section_title(doc, 'Certifications')
        for cert in certs:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(0)
            
            run_name = p.add_run(cert['name'])
            run_name.bold = True
            
            run_prov = p.add_run(f" – {cert['provider']} – ")
            run_prov.italic = True
            
            self.add_hyperlink(p, cert['url'], "Link", "0000FF", True, False)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def _add_experience(self, doc, jobs):
        if not jobs: return
        self._add_section_title(doc, 'Experience')
        for job in jobs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), alignment=WD_TAB_ALIGNMENT.RIGHT)
            
            run1 = p.add_run(job['title'])
            run1.bold = True
            
            run2 = p.add_run(f"\t{job['date']}")
            
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(1)
            p2.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), alignment=WD_TAB_ALIGNMENT.RIGHT)
            
            run3 = p2.add_run(job['company'])
            run3.italic = True
            
            run4 = p2.add_run(f"\t{job['location']}")
            run4.italic = True

            for bullet in job['bullets']:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(0) 
                run = p.add_run(bullet)
                run.font.size = Pt(10)
                
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def _add_projects(self, doc, projects):
        if not projects: return
        self._add_section_title(doc, 'Projects')
        for proj in projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), alignment=WD_TAB_ALIGNMENT.RIGHT)
            
            run_name = p.add_run(proj['name'])
            run_name.bold = True
            
            p.add_run(" | ")
            
            run_stack = p.add_run(proj['stack'])
            run_stack.italic = True
            
            p.add_run(f"\t{proj['date']}")

            for bullet in proj['bullets']:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(bullet)
                run.font.size = Pt(10)
                
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def _add_skills(self, doc, skills):
        if not skills: return
        self._add_section_title(doc, 'Technical Skills')
        for category, items in skills.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run_cat = p.add_run(f"{category}: ")
            run_cat.bold = True
            p.add_run(items)

    def _add_row_two_cols(self, doc, col1_text, col1_bold, col1_italic, col2_text, col2_bold, col2_italic, space_after=Pt(0)):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = space_after
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), alignment=WD_TAB_ALIGNMENT.RIGHT)
        
        run1 = p.add_run(col1_text)
        run1.bold = col1_bold
        run1.italic = col1_italic
        
        run2 = p.add_run(f"\t{col2_text}")
        run2.bold = col2_bold
        run2.italic = col2_italic
